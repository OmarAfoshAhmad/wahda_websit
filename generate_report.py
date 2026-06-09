import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    
    # 1. Add <w:bidi/> to paragraph properties
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    # 2. Add <w:jc w:val="both"/> to justify text
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

def set_run_rtl(run):
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        set_run_rtl(run)
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Text formatting
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    run.font.size = Pt(13)
    run.bold = bold
    set_run_rtl(run)
    
    return p

def main():
    doc = Document()
    
    # Set document-wide style to Arabic Language and RTL orientation
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(13)
    normal_font.rtl = True
    
    # Add Logo
    logo_path = r"c:\Users\Omar\waad_temp_website\public\logo.png"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(2.0))
    
    # Title
    title = doc.add_heading('التقرير التقني والفني لمنظومة الواحة للرعاية الصحية', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    for run in title.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        set_run_rtl(run)
    
    doc.add_paragraph() # Spacing
    
    # Section 1
    add_heading(doc, 'مقدمة عن المنظومة', 1)
    add_paragraph(doc, 'منظومة الواحة للرعاية الصحية تمثل منصة رقمية متطورة وحديثة لإدارة الطرف الثالث في قطاع التأمين الصحي والرعاية الطبية. صُممت المنظومة لتكون حلقة الوصل الفعّالة والموثوقة بين المستفيدين وشركات التأمين والمرافق الطبية المتعاقدة معها، مما يوفر بيئة متكاملة وآمنة وشفافة لإدارة السقوف المالية ومعالجة المطالبات الطبية بالإضافة إلى تتبع الحركات والخصومات بشكل لحظي ودقيق.')
    
    # Section 2
    add_heading(doc, 'البنية التقنية المتطورة', 1)
    add_paragraph(doc, 'تعتمد منظومة الواحة على أحدث التقنيات البرمجية المتوفرة لضمان الأداء العالي والأمان والموثوقية المطلقة وقابلية التوسع السلسة. تم بناء واجهة المستخدم الأمامية باستخدام أطر عمل حديثة ومتقدمة توفر تجربة استخدام تفاعلية ومتجاوبة عبر مختلف الأجهزة والشاشات. أما الواجهة الخلفية فتعتمد على معمارية معالجة مركزية تستخدم بيئة تشغيل آمنة توفر سرعة استجابة استثنائية وأماناً متقدماً في تنفيذ العمليات الخادمة المعقدة. وفيما يخص تخزين البيانات، تعتمد المنظومة على قواعد بيانات علائقية متينة وموثوقة لضمان تكامل البيانات وسرعة الاستعلامات وتأمين السجلات الحساسة ضد أي فقدان محتمل.')
    
    # Section 3
    add_heading(doc, 'الميزات والأقسام الرئيسية', 1)
    
    add_heading(doc, 'إدارة المستفيدين والمحافظ الطبية', 2)
    add_paragraph(doc, 'تقدم المنظومة آليات ترقيم وتسجيل ذكية ومبتكرة للبطاقات تضمن التفرد وعدم تكرار السجلات بشكل قاطع. وتتميز بوجود محرك استيراد وتصدير فائق السرعة يتيح استيعاب آلاف السجلات دفعة واحدة مع مراعاة كافة ضوابط البيانات. يتيح هذا القسم تقديم إدارة شاملة ودقيقة للأرصدة الكلية والمتبقية للمستفيد، مع ربط أوتوماتيكي ومباشر بأسقف وتغطيات شركات التأمين المعتمدة لضمان التحديث المستمر للأرصدة.')
    
    add_heading(doc, 'هندسة الحركات والمطالبات', 2)
    add_paragraph(doc, 'يتميز النظام بمحرك تسوية حركات لحظي يقوم بالخصم المباشر من أرصدة المستفيدين بمجرد إجراء الخدمات الطبية في المرافق المعتمدة. كما يوفر نظام استرداد وتسوية متطور يعيد الأرصدة للمستفيدين أوتوماتيكياً وبدقة متناهية عند إلغاء أو تصحيح أي حركة. إلى جانب ذلك، تتعامل المنظومة مع الخصومات المخصصة مثل خدمات الأسنان والخدمات الدوائية من خلال حسابات استهلاكية منفصلة ومستقلة تعكس السقوف الفرعية بدقة تامة.')
    
    add_heading(doc, 'بوابة المستفيد الإلكترونية', 2)
    add_paragraph(doc, 'صُممت بوابة المستفيد لتكون المساحة الآمنة والخاصة للمريض، حيث تتيح الدخول حصرياً عبر روابط مؤمنة مقترنة برموز التحقق المؤقتة لضمان أقصى درجات الخصوصية للبيانات الطبية والمالية. تمكن البوابة المستفيد من تعيين رموز وصول شخصية واستعراض رصيده الفعلي وسجل حركاته التفصيلي وتتبع المبالغ المستهلكة بشفافية تامة وبدون تعقيدات.')
    
    add_heading(doc, 'الصيانة الدورية وتدقيق البيانات', 2)
    add_paragraph(doc, 'لضمان نزاهة العمليات، تحتوي المنظومة على سجل تدقيق شامل وأمني يتتبع كل عملية إضافة أو تعديل أو حذف تحدث داخل المنصة، مع توثيق اسم المستخدم وتوقيت الإجراء. علاوة على ذلك، توفر المنظومة خوارزميات مدمجة لمعالجة السجلات المكررة والديون المتجاوزة للسقوف بفاعلية عالية، وتضمن استمرارية الأعمال بفضل نظام النسخ الاحتياطي التلقائي للبيانات.')
    
    # Section 4
    add_heading(doc, 'مستويات الأمان والحماية', 1)
    add_paragraph(doc, 'تطبق منظومة الواحة أعلى معايير أمن المعلومات العالمية من خلال تشفير قنوات الاتصال واستخدام صلاحيات الوصول الدقيقة والمبنية على الأدوار والتي تمنع بشكل صارم أي وصول غير مصرح به للوظائف الإدارية أو السجلات الطبية. يتم فرض جدران حماية متقدمة على مستوى الجلسات وتأكيد هوية المستفيدين عبر بروتوكولات التحقق المستمر لضمان أن كل حركة وكل استعلام يتم تحت توثيق كامل ورقابة سيبرانية مشددة.')
    
    # Section 5
    add_heading(doc, 'الخلاصة', 1)
    add_paragraph(doc, 'تمثل منظومة الواحة حلاً تقنياً وحوكمياً متكاملاً يتجاوز كونه مجرد تطبيق لإدارة البيانات، ليصبح ركيزة أساسية في التحول الرقمي والحوكمة الطبية وإدارة النفقات التأمينية. إنها تتميز بمرونتها الهائلة في التكيف مع ديناميكيات سياسات التأمين وتوفير الشفافية المطلقة والسرعة الموثوقة لجميع الأطراف المعنية في منظومة الرعاية الصحية.')
    
    # Save the file to desktop
    output_path = os.path.join(os.path.expanduser("~"), "Desktop", "تقرير_منظومة_الواحة_التقني_المحدث.docx")
    doc.save(output_path)
    print("Report generated successfully at:", output_path)

if __name__ == "__main__":
    main()
